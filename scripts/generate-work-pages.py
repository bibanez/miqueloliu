#!/usr/bin/env python3
"""Generate static work-detail pages from the working Word documents.

The Word files in ``documentacio/obres`` are editorial sources and remain
gitignored. This script produces the publishable HTML in ``obres/`` and a
small metadata index in ``js/data/work-info.js``. The metadata lets the
catalogue know which notes exist and which of their source languages overlap
with the website's Catalan/Spanish/English UI languages.

The source files are not consistently labelled internally, so the language
ranges are explicit. This prevents extra French, German, or Euskera material
from being silently discarded and makes adding another source language a
small, reviewable manifest change.

Run with the bundled Python runtime:

    python3 scripts/generate-work-pages.py
"""

from __future__ import annotations

import argparse
import html
import json
import re
import shutil
import subprocess
import tempfile
import unicodedata
from dataclasses import dataclass
from pathlib import Path

from docx import Document


SITE_LANGUAGES = ("ca", "es", "en")
LANGUAGE_ORDER = ("ca", "es", "en", "fr", "de", "eu")
LANGUAGE_LABELS = {
    "ca": "Català",
    "es": "Castellano",
    "en": "English",
    "fr": "Français",
    "de": "Deutsch",
    "eu": "Euskara",
}
LANGUAGE_CODES = {
    "ca": "CA",
    "es": "ES",
    "en": "EN",
    "fr": "FR",
    "de": "DE",
    "eu": "EU",
}
NOTE_HEADINGS = {
    "nota de programa",
    "programme note",
    "program note",
    "programa note",
    "a propos du programme",
    "preface",
    "note",
}


WORKS = [
    {
        "id": "pluja",
        "source": "Pluja Nota de programa i poemes Cat -Cast-Eusk-Engl-Deutsch-Francès.docx",
        "titles": {
            "ca": "Pluja",
            "es": "Pluja (Lluvia)",
            "en": "Pluja (Rain)",
            "eu": "Euri",
            "de": "Regen",
            "fr": "Pluie",
        },
        "segments": {
            "ca": (0, 49),
            "es": (49, 107),
            "eu": (107, 145),
            "en": (145, 199),
            "de": (199, 236),
            "fr": (236, 268),
        },
    },
    {
        "id": "in-luce-praesenti",
        "source": "In luce Praesenti tres idiomes.docx",
        "titles": {"ca": "In luce praesenti", "es": "In luce praesenti", "en": "In luce praesenti"},
        "tables": {"ca": 0, "es": 1, "en": 2},
    },
    {
        "id": "allegories-de-tardor-orquestra",
        "source": "Al·legories de tardor (orquestra) tres idiomes.docx",
        "titles": {
            "ca": "Al·legories de tardor – Microludis per a orquestra",
            "es": "Al·legories de tardor – Microludis per a orquestra (Alegorías de otoño – Microludios para orquesta)",
            "en": "Al·legories de tardor – Microludis per a orquestra (Autumn allegories – Microludes for orchestra)",
        },
        "segments": {"ca": (0, 40), "es": (40, 58), "en": (58, 74)},
    },
    {
        "id": "un-simple-aleteig",
        "source": "Un simple aleteig... Nota de programa definitiu CAT CAST ENG.docx",
        "titles": {
            "ca": "Un simple aleteig...",
            "es": "Un simple aleteig... (Un simple aleteo...)",
            "en": "Un simple aleteig... (A simple flutter...)",
        },
        "segments": {"ca": (0, 14), "es": (14, 27), "en": (27, 39)},
    },
    {
        "id": "set-cants",
        "source": "Set cants Nota de programa i poemes (CAT-FR-ENG).docx",
        "titles": {"ca": "Set cants", "es": "Set cants (Siete cantos)", "en": "Set cants (Seven songs)"},
        "segments": {"ca": (0, 124), "es": (124, 128), "fr": (140, 270), "en": (270, 383)},
    },
    {
        "id": "angel-terrible",
        "source": "L'àngel terrible (tres idiomes).doc",
        "titles": {
            "ca": "L’àngel terrible",
            "es": "L’àngel terrible (El ángel terrible)",
            "en": "L’àngel terrible (The terrible angel)",
        },
        "segments": {"ca": (0, 32), "es": (32, 57), "en": (57, 83)},
    },
    {
        "id": "jardins-del-silenci",
        "source": "Jardins del silenci nota de programa tres idiomes.doc",
        "titles": {
            "ca": "Jardins del silenci",
            "es": "Jardins del silenci (Jardines del silencio)",
            "en": "Jardins del silenci (Gardens of silence)",
        },
        "segments": {"ca": (0, 26), "es": (26, 51), "en": (51, 64)},
    },
    {
        "id": "ofrena-a-ant-negre",
        "source": "Ofrena a Ant Negre tres idiomes.docx",
        "titles": {
            "ca": "Ofrena a Ant Negre",
            "es": "Ofrena a Ant Negre (Ofrenda a Alce Negro)",
            "en": "Ofrena a Ant Negre (Offering to Black Elk)",
        },
        "segments": {"ca": (0, 15), "es": (15, 29), "en": (29, 40)},
    },
    {
        "id": "inabastable-perfil",
        "source": "L'inabastable perfil  tres idiomes.doc",
        "titles": {
            "ca": "L’inabastable perfil – Homenatge E. Chillida",
            "es": "L’inabastable perfil – Homenatge E. Chillida (El inalcanzable perfil – Homenaje a E. Chillida)",
            "en": "L’inabastable perfil – Homenatge E. Chillida (The unattainable profile – Homage to E. Chillida)",
        },
        "segments": {"ca": (0, 10), "es": (10, 21), "en": (21, 35)},
    },
    {
        "id": "set-extractes-petit-princep",
        "source": "Set extractes El Petit princep  comentari obra CAT-CAST-ENG.doc",
        "titles": {
            "ca": "Set extractes de El Petit Príncep",
            "es": "Set extractes de El Petit Príncep (Siete extractos de El principito)",
            "en": "Set extractes de El Petit Príncep (Seven extracts from The Little Prince)",
        },
        "segments": {"ca": (0, 5), "es": (5, 10), "en": (10, 15)},
    },
    {
        "id": "lament-de-pluja",
        "source": "Lament de pluja  CAT-CAST-ENG.docx",
        "titles": {
            "ca": "Lament de pluja",
            "es": "Lament de pluja (Lamento de lluvia)",
            "en": "Lament de pluja (Lament of rain)",
        },
        "segments": {"ca": (0, 2), "es": (2, 4), "en": (4, 7)},
    },
    {
        "id": "llibre-dhores",
        "source": "Llibre d'hores_tres idiomes.docx",
        "titles": {
            "ca": "Llibre d’hores. Cicle complert de Preludis per a piano, en quatre quaderns",
            "es": "Llibre d’hores (Libro de horas). Ciclo completo de Preludios para piano, en cuatro cuadernos",
            "en": "Llibre d’hores (Book of hours). Complete set of the Preludes for piano, in four volumes",
        },
        "segments": {"ca": (0, 56), "es": (56, 117), "en": (117, 166)},
    },
    {
        "id": "cantic",
        "source": "Càntic flauta sola Nota de programa tres idiomes.docx",
        "titles": {"ca": "Càntic", "es": "Càntic (Cántico)", "en": "Càntic (Canticle)"},
        "segments": {"ca": (0, 7), "es": (7, 15), "en": (15, 26)},
    },
    {
        "id": "fondres",
        "source": "Fondre's tres idiomes.doc",
        "titles": {"ca": "Fondre’s", "es": "Fondre’s (Fundirse)", "en": "Fondre’s (Melting)"},
        "segments": {"ca": (0, 27), "es": (27, 54), "en": (54, 74)},
    },
    {
        "id": "despertar",
        "source": "Despertar tres idiomes.docx",
        "titles": {"ca": "Despertar", "es": "Despertar", "en": "Despertar (Awakening)"},
        "segments": {"ca": (0, 6), "es": (6, 11), "en": (11, 16)},
    },
    {
        "id": "allegories-de-tardor-piano",
        "source": "Al·legories de tardor (Piano) 4 idiomes.docx",
        "titles": {
            "ca": "Al·legories de tardor – Microludis per a piano",
            "es": "Al·legories de tardor – Microludis per a piano (Alegorías de otoño – Microludios para piano)",
            "en": "Al·legories de tardor – Microludis per a piano (Autumn allegories – Microludes for piano)",
            "de": "Mikroludien – Herbst Allegorien",
        },
        "segments": {"ca": (0, 18), "es": (18, 33), "en": (33, 48), "de": (48, 62)},
    },
    {
        "id": "sant-josep-va-a-buscar-foc",
        "source": "St. Josep va a buscar foc   (dedicatòria 3 idiomes)).docx",
        "titles": {"ca": "Sant Josep va a buscar foc", "es": "Sant Josep va a buscar foc", "en": "Sant Josep va a buscar foc"},
        "segments": {"ca": (0, 3), "es": (3, 6), "en": (6, 11)},
    },
    {
        "id": "jo-laimo",
        "source": "Jo l'aimo Nota de programa CAT- CAST.docx",
        "titles": {"ca": "Jo l’aimo", "es": "Jo l’aimo (Yo la amo)", "en": "Jo l’aimo (I love her)"},
        "segments": {"ca": (0, 6), "es": (6, 10)},
    },
]


@dataclass
class Block:
    text: str
    style: str = "Normal"
    blank: bool = False
    poem: bool = False


def normalize(value: str) -> str:
    value = unicodedata.normalize("NFKC", value).replace("\u00a0", " ")
    return " ".join(value.split()).strip().lower()


def is_zero_length(value) -> bool:
    return value is not None and (value == 0 or getattr(value, "pt", None) == 0)


def clean_text(value: str) -> str:
    return value.replace("\x00", "").replace("\r", "").rstrip()


def paragraph_to_block(paragraph) -> Block:
    text = clean_text(paragraph.text)
    return Block(
        text=text,
        style=paragraph.style.name if paragraph.style else "Normal",
        blank=not text.strip(),
        # Word uses zero spacing for both verse lines and some compact prose.
        # Long sentences should stay prose even if their source uses that
        # spacing setting.
        poem=is_zero_length(paragraph.paragraph_format.space_after) and len(text.strip()) <= 120,
    )


def convert_legacy_doc(path: Path, temp_dir: Path) -> Path:
    target = temp_dir / f"{path.stem}.docx"
    textutil = shutil.which("textutil")
    if textutil:
        subprocess.run(
            [textutil, "-convert", "docx", "-output", str(target), str(path)],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
        )
        return target

    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if soffice:
        subprocess.run(
            [soffice, "--headless", "--convert-to", "docx", "--outdir", str(temp_dir), str(path)],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
        )
        if target.exists():
            return target

    raise RuntimeError(f"Cannot convert legacy Word document: {path}")


def read_source(path: Path, temp_dir: Path):
    converted = convert_legacy_doc(path, temp_dir) if path.suffix.lower() == ".doc" else path
    document = Document(converted)
    paragraphs = [paragraph_to_block(p) for p in document.paragraphs]
    tables = []
    for table in document.tables:
        rows = []
        for row in table.rows:
            for cell in row.cells:
                rows.extend(paragraph_to_block(p) for p in cell.paragraphs)
        tables.append(rows)
    return paragraphs, tables


def inline_text(value: str) -> str:
    return html.escape(value, quote=False).replace("\n", "<br>")


def is_heading(block: Block) -> int | None:
    text = normalize(block.text)
    if not text:
        return None
    if block.style.lower().startswith("heading"):
        return 3
    if text in NOTE_HEADINGS:
        return 2
    if re.fullmatch(r"[ivxlcdm]+\.?", text):
        return 3
    if re.match(r"^(preludi|preludio|prelude)\s+[ivxlcdm]+\b", text):
        return 3
    if re.match(r"^(primer|segon|tercer|quart|first|second|third|fourth)\s+(quadern|cuaderno|volume|volum)", text):
        return 3
    return None


def is_list_style(block: Block) -> bool:
    return block.style.lower().startswith("list")


def render_blocks(blocks: list[Block], title_values: list[str]) -> str:
    title_norms = [normalize(value) for value in title_values]
    output = []
    in_poem = False

    def close_poem():
        nonlocal in_poem
        if in_poem:
            output.append("</div>")
            in_poem = False

    def is_duplicate_title(text: str) -> bool:
        normalized = normalize(text)
        return any(
            normalized == candidate
            or (len(normalized) > 8 and candidate.startswith(normalized))
            for candidate in title_norms
        )

    index = 0
    while index < len(blocks):
        block = blocks[index]
        if block.blank:
            if in_poem:
                next_index = index + 1
                while next_index < len(blocks) and blocks[next_index].blank:
                    next_index += 1
                if next_index < len(blocks) and blocks[next_index].poem:
                    if not output or not output[-1].startswith('<div class="work-info-stanza-break"'):
                        output.append('<div class="work-info-stanza-break" aria-hidden="true"></div>')
                else:
                    close_poem()
            index += 1
            continue

        if is_list_style(block):
            close_poem()
            items = []
            while index < len(blocks) and (is_list_style(blocks[index]) or blocks[index].blank):
                if not blocks[index].blank:
                    items.append(f"<li>{inline_text(blocks[index].text.strip())}</li>")
                index += 1
            if items:
                output.append('<ul class="work-info-list">' + "".join(items) + "</ul>")
            continue

        level = is_heading(block)
        if level:
            close_poem()
            if not (level == 2 and is_duplicate_title(block.text)):
                output.append(f"<h{level}>{inline_text(block.text.strip())}</h{level}>")
            index += 1
            continue

        if block.poem:
            if not in_poem:
                output.append('<div class="work-info-poem">')
                in_poem = True
            output.append(f'<p class="work-info-poem-line">{inline_text(block.text)}</p>')
        else:
            close_poem()
            output.append(f"<p>{inline_text(block.text)}</p>")
        index += 1

    close_poem()
    return "\n          ".join(output)


def source_languages(entry) -> list[str]:
    values = list(entry.get("segments", {}).keys()) + list(entry.get("tables", {}).keys())
    return sorted(set(values), key=lambda value: LANGUAGE_ORDER.index(value) if value in LANGUAGE_ORDER else len(LANGUAGE_ORDER))


def segment_blocks(entry, lang: str, paragraphs, tables):
    if lang in entry.get("tables", {}):
        return tables[entry["tables"][lang]]
    if lang in entry.get("segments", {}):
        start, end = entry["segments"][lang]
        return paragraphs[start:end]
    raise KeyError(f"No source range configured for {entry['id']} / {lang}")


def render_page(entry, rendered, languages):
    title_ca = entry["titles"].get("ca") or entry["titles"].get(languages[0], entry["id"])
    title_html = "\n          ".join(
        f'<div class="work-detail-lang-block" data-detail-lang="{lang}"><h1>{inline_text(entry["titles"].get(lang, title_ca))}</h1></div>'
        for lang in languages
    )
    tabs_html = "\n            ".join(
        f'<button class="work-detail-language-tab" type="button" role="tab" data-work-language="{lang}" aria-label="{html.escape(LANGUAGE_LABELS.get(lang, lang.upper()), quote=True)}" aria-selected="false">{html.escape(LANGUAGE_CODES.get(lang, lang[:2].upper()))}</button>'
        for lang in languages
    )
    content_html = "\n          ".join(
        f'<div class="work-detail-lang-block" data-detail-lang="{lang}"><div class="work-info">{rendered[lang]}</div></div>'
        for lang in languages
    )
    description = html.escape(f"Nota de programa de {title_ca}.", quote=True)
    return f'''<!DOCTYPE html>
<html lang="ca">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <meta name="description" content="{description}">
  <title>Miquel Oliu — {html.escape(title_ca, quote=False)}</title>
  <base href="../">
  <script>
    try {{
      var storedLang = localStorage.getItem('miqueloliu_lang');
      if (storedLang === 'es' || storedLang === 'en') document.documentElement.lang = storedLang;
    }} catch (e) {{}}
  </script>
  <link rel="preconnect" href="https://fonts.googleapis.com">
  <link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
  <link href="https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500&family=Playfair+Display:ital,wght@0,400;0,500;1,400&display=swap" rel="stylesheet">
  <link rel="stylesheet" href="css/style.css">
</head>
<body>
  <div id="nav-container"></div>

  <main data-page="work-detail" data-work-id="{entry["id"]}">
    <article class="work-detail-layout">
      <div class="page-header work-detail-header">
        <p class="work-detail-back"><a href="catalogue.html#work-{entry["id"]}" data-i18n="works.detail.back"></a></p>
        {title_html}
        <nav class="work-detail-language-nav" aria-label="Programme note languages">
          <span class="work-detail-language-label" data-i18n="works.detail.language.label"></span>
          <div class="work-detail-language-tabs" role="tablist">
            {tabs_html}
          </div>
        </nav>
      </div>
      <div class="work-detail-content">
        {content_html}
      </div>
      <p class="work-detail-back work-detail-back-bottom"><a href="catalogue.html#work-{entry["id"]}" data-i18n="works.detail.back"></a></p>
    </article>
  </main>

  <div id="footer-container"></div>

  <script src="js/translations.js"></script>
  <script src="js/i18n.js"></script>
  <script src="js/components.js"></script>
  <script src="js/work-detail.js"></script>
  <script src="js/main.js"></script>
</body>
</html>
'''


def render_metadata(entries):
    metadata = {}
    for entry in entries:
        languages = source_languages(entry)
        metadata[entry["id"]] = {
            "href": f'obres/{entry["id"]}.html',
            "languages": languages,
            "siteLanguages": [lang for lang in languages if lang in SITE_LANGUAGES],
            "hasSiteLanguage": any(lang in SITE_LANGUAGES for lang in languages),
        }
    return "/* Generated by scripts/generate-work-pages.py. */\nconst WORK_INFO = " + json.dumps(metadata, ensure_ascii=False, indent=2) + ";\n"


def main():
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--source-dir", type=Path, default=Path("documentacio/obres"))
    parser.add_argument("--output-dir", type=Path, default=Path("obres"))
    parser.add_argument("--metadata-output", type=Path, default=Path("js/data/work-info.js"))
    args = parser.parse_args()

    args.output_dir.mkdir(parents=True, exist_ok=True)
    args.metadata_output.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="miquel-works-") as temp:
        temp_dir = Path(temp)
        for entry in WORKS:
            source = args.source_dir / entry["source"]
            if not source.exists():
                raise FileNotFoundError(source)
            paragraphs, tables = read_source(source, temp_dir)
            languages = source_languages(entry)
            rendered = {
                lang: render_blocks(segment_blocks(entry, lang, paragraphs, tables), list(entry["titles"].values()))
                for lang in languages
            }
            destination = args.output_dir / f'{entry["id"]}.html'
            destination.write_text(render_page(entry, rendered, languages), encoding="utf-8")
            print(f"generated {destination} ({', '.join(languages)})")

    args.metadata_output.write_text(render_metadata(WORKS), encoding="utf-8")
    print(f"generated {args.metadata_output}")


if __name__ == "__main__":
    main()
